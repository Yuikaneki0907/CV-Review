import io
import os
import re
import tempfile
import traceback
from datetime import datetime
from typing import List, Literal
from uuid import UUID, uuid4

from docx import Document
from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile
from fastapi.responses import Response, StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.application.dto.requests import (
    ChatContextRequest,
    GeneratedCVExportPreviewRequest,
    ChatSessionMessagesRequest,
    GenerateCVRequest,
    GeneratedCVUpdateRequest,
)
from app.application.dto.responses import (
    ChatContextResponse,
    ChatSessionResponse,
    GeneratedCVListResponse,
    GeneratedCVResponse,
    GeneratedCVVersionResponse,
)
from app.application.use_cases.edit_generated_cv import EditGeneratedCVUseCase
from app.application.use_cases.generate_cv import GenerateCVUseCase
from app.application.use_cases.import_generated_cv import ImportGeneratedCVUseCase
from app.application.use_cases.chat_cv import ChatCVUseCase
from app.infrastructure.ai import ai_service_factory
from app.infrastructure.database.session import get_db_session
from app.infrastructure.database.repositories.generated_cv_repository import GeneratedCVRepository
from app.infrastructure.database.repositories.user_repository import UserRepository
from app.infrastructure.database.models import ChatSessionModel, GeneratedCVModel
from app.infrastructure.file_parsers.parsers import get_parser
from app.infrastructure.file_parsers.import_pipeline import (
    build_import_preview_payload,
    convert_pdf_to_docx,
)
from app.infrastructure.file_parsers.upload_validation import read_and_validate_upload
from app.presentation.auth_routes import get_current_user_id
from app.logger import get_logger
from sqlalchemy import and_, func, select, update

logger = get_logger("app.presentation.generated_cv_routes")

router = APIRouter(prefix="/generated-cvs", tags=["Generated CVs"])
MAX_IMPORTED_CV_SIZE_MB = 5


def _strip_markdown_inline(text: str) -> str:
    content = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1 (\2)", text)
    content = content.replace("**", "").replace("*", "").replace("`", "")
    return content.strip()


def _markdown_to_docx_bytes(markdown_text: str) -> bytes:
    doc = Document()

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            heading_level = min(len(heading_match.group(1)), 4)
            doc.add_heading(_strip_markdown_inline(heading_match.group(2)), level=heading_level)
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", line) or re.match(r"^\d+\.\s+(.*)$", line)
        if bullet_match:
            doc.add_paragraph(_strip_markdown_inline(bullet_match.group(1)), style="List Bullet")
            continue

        doc.add_paragraph(_strip_markdown_inline(line))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _get_generated_content_payload(cv_entity) -> tuple[str, str]:
    content_data = cv_entity.generated_content if isinstance(cv_entity.generated_content, dict) else {}
    output_format = content_data.get("format")

    if output_format not in {"markdown", "docx"}:
        if isinstance(content_data.get("markdown"), str):
            output_format = "markdown"
        else:
            output_format = "markdown"

    content = (
        content_data.get("content")
        or content_data.get("markdown")
        or ""
    )
    return output_format, content


def _build_export_filename(cv_entity, ext: str) -> str:
    job_title = (cv_entity.base_profile_data or {}).get("job_title") or "generated_cv"
    normalized = re.sub(r"[^a-zA-Z0-9_-]+", "_", str(job_title).strip().lower()).strip("_")
    if not normalized:
        normalized = "generated_cv"
    return f"{normalized[:60]}.{ext}"



GENERIC_CHAT_TITLES = {
    "cv từ chatbot",
    "cv tu chatbot",
    "cv đã tạo",
    "cv da tao",
    "generated_cv",
    "generated cv",
    "cv unnamed",
    "cuộc trò chuyện mới",
    "cuoc tro chuyen moi",
}

ROLE_TITLE_PATTERNS = [
    r"\b(?:ai|machine learning|ml|data|backend|frontend|fullstack|full-stack|software|web|mobile|devops|qa|tester|business analyst|ba)\s+(?:engineer|developer|intern|fresher|analyst|specialist)\b",
    r"\b(?:intern|fresher|junior|middle|mid-level|senior)\s+(?:ai|machine learning|ml|data|backend|frontend|fullstack|full-stack|software|web|mobile|devops|qa|tester|developer|engineer|analyst)\b",
    r"\b(?:java|python|react|node(?:\.js)?|php|\.net|c#|golang|android|ios)\s+(?:developer|engineer|intern|fresher)\b",
]

ROLE_AFTER_MARKERS = [
    r"(?:role|vị trí|vi tri|ứng tuyển|ung tuyen|apply|cho anh cv|tạo cv|tao cv|viết cv|viet cv)\s+(?:là|la|cho|role|vị trí|vi tri)?\s*[:\-]?\s*([a-zA-ZÀ-ỹ0-9+#./ -]{3,60})",
]

def _normalize_title_text(text: str, limit: int = 72) -> str:
    value = str(text or "").strip()
    if not value:
        return ""
    value = re.sub(r"```.*?```", " ", value, flags=re.DOTALL)
    value = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", value)
    value = re.sub(r"[*_`>#]+", " ", value)
    value = re.sub(r"https?://\S+", " ", value)
    value = re.sub(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", " ", value)
    value = re.sub(r"\s+", " ", value).strip(" -:.,;")
    return value[:limit].rstrip(" -:.,;")

def _is_generic_title(text: str) -> bool:
    normalized = _normalize_title_text(text, 120).lower()
    return not normalized or normalized in GENERIC_CHAT_TITLES

def _title_case_role(role: str) -> str:
    cleaned = _normalize_title_text(role, 64)
    cleaned = re.sub(
        r"\b(mà|ma|và|va|với|voi|theo|cho|của|cua|này|nay|đó|do|nó|no|match|khớp|khop)\b.*$",
        "",
        cleaned,
        flags=re.IGNORECASE,
    ).strip(" -:.,;")
    if not cleaned:
        return ""
    aliases = {
        "aie": "AI Engineer",
        "ai engineer": "AI Engineer",
        "ai intern": "AI Intern",
        "ai/ml intern": "AI/ML Intern",
        "ml intern": "ML Intern",
        "backend intern": "Backend Intern",
        "frontend intern": "Frontend Intern",
        "software engineer": "Software Engineer",
        "software engineer intern": "Software Engineer Intern",
        "data analyst": "Data Analyst",
        "data analyst intern": "Data Analyst Intern",
    }
    lowered = cleaned.lower()
    if lowered in aliases:
        return aliases[lowered]
    aie_match = re.fullmatch(r"(?:(intern|fresher|junior|senior)\s+)?aie", lowered)
    if aie_match:
        level = aie_match.group(1)
        return f"{level[:1].upper() + level[1:]} AI Engineer" if level else "AI Engineer"
    words = []
    for word in cleaned.split():
        lower_word = word.lower()
        if lower_word in {"ai", "ml", "qa", "ba", "ui", "ux", "jd", "cv"}:
            words.append(lower_word.upper())
        elif lower_word in {"c#", ".net", "node.js"}:
            words.append(word)
        else:
            words.append(word[:1].upper() + word[1:])
    return " ".join(words)

def _extract_role_title(text: str) -> str:
    normalized = re.sub(r"\s+", " ", str(text or "")).strip()
    if not normalized:
        return ""

    for pattern in ROLE_TITLE_PATTERNS:
        match = re.search(pattern, normalized, flags=re.IGNORECASE)
        if match:
            return _title_case_role(match.group(0))

    for pattern in ROLE_AFTER_MARKERS:
        match = re.search(pattern, normalized, flags=re.IGNORECASE)
        if not match:
            continue
        candidate = _title_case_role(match.group(1))
        if candidate and len(candidate.split()) <= 6:
            return candidate

    return ""

def _extract_title_from_cv_markdown(markdown_text: str) -> str:
    text = str(markdown_text or "")
    if not text.strip():
        return ""

    heading_match = re.search(r"(?m)^#\s+(.+?)\s*$", text)
    if heading_match:
        heading = _normalize_title_text(heading_match.group(1), 72)
        if heading and not re.search(r"\[[^\]]+\]|tên của bạn|your name", heading, flags=re.IGNORECASE):
            return heading

    return _extract_role_title(text[:2500])

def _build_first_query_title(chat_history) -> str | None:
    if not isinstance(chat_history, list):
        return None

    for message in chat_history:
        if not isinstance(message, dict) or message.get("role") != "user":
            continue
        content = str(message.get("content") or "").strip()
        if not content:
            continue

        lines = [
            _normalize_title_text(line, 180)
            for line in content.splitlines()
            if _normalize_title_text(line, 180)
        ]
        first_line = lines[0] if lines else _normalize_title_text(content, 180)
        return first_line or None

    return None

def _build_chat_title_from_messages(chat_history) -> str | None:
    if not isinstance(chat_history, list):
        return None

    first_query_title = _build_first_query_title(chat_history)
    if first_query_title:
        return first_query_title

    user_messages = [
        str(message.get("content") or "").strip()
        for message in chat_history
        if isinstance(message, dict)
        and message.get("role") == "user"
        and str(message.get("content") or "").strip()
    ]
    if not user_messages:
        return None

    combined = "\n".join(user_messages)
    lowered = combined.lower()

    if "jd" in lowered and "cv" in lowered and any(token in lowered for token in ["match", "khớp", "khop", "tương ứng", "tuong ung"]):
        return "Tạo JD và CV khớp nhau"
    if any(token in lowered for token in ["phân tích cv", "phan tich cv", "phân tích tài liệu", "phan tich tai lieu"]) and any(token in lowered for token in ["jd", "job", "mô tả", "mo ta"]):
        return "Phân tích CV theo JD"
    if "tài liệu đính kèm" in lowered or "tai lieu dinh kem" in lowered:
        return "Xử lý tài liệu đính kèm"
    if any(token in lowered for token in ["mẫu cv", "mau cv", "template cv", "các mẫu", "cac mau"]):
        return "Tư vấn mẫu CV"

    role_title = _extract_role_title(combined)
    if role_title and any(token in lowered for token in ["cv", "resume", "hồ sơ", "ho so", "ứng tuyển", "ung tuyen"]):
        return f"CV {role_title}" if not role_title.lower().startswith("cv ") else role_title

    if any(token in lowered for token in ["gen cv", "tạo cv", "tao cv", "viết cv", "viet cv"]):
        return "Tạo CV bằng AI"
    if any(token in lowered for token in ["sửa cv", "sua cv", "chỉnh cv", "chinh cv", "rewrite", "cập nhật", "cap nhat"]):
        return "Chỉnh sửa CV"
    if any(token in lowered for token in ["tôi tên", "toi ten", "anh tên", "em tên", "học trường", "hoc truong", "chuyên ngành", "chuyen nganh", "kinh nghiệm", "kinh nghiem", "dự án", "du an"]):
        return "Hoàn thiện thông tin CV"

    fallback = _normalize_title_text(user_messages[0], 48)
    if fallback and len(fallback.split()) <= 8:
        return fallback
    return None

def _build_chat_title(chat_history, latest_cv: GeneratedCVModel | None = None) -> str | None:
    first_query_title = _build_first_query_title(chat_history)
    if first_query_title:
        return first_query_title

    if latest_cv:
        profile = latest_cv.base_profile_data if isinstance(latest_cv.base_profile_data, dict) else {}
        source_type, source_filename = _resolve_source_info(profile)
        if source_type == "uploaded_cv" and source_filename:
            return _normalize_title_text(f"CV tải lên: {source_filename}", 72)

        job_title = profile.get("job_title")
        if isinstance(job_title, str) and not _is_generic_title(job_title):
            normalized_job = _title_case_role(job_title)
            return _normalize_title_text(
                normalized_job if normalized_job.lower().startswith("cv ") else f"CV {normalized_job}",
                72,
            )

        target_role = _extract_role_title(latest_cv.target_jd_text or "")
        if target_role:
            return _normalize_title_text(f"CV {target_role}", 72)

        content = latest_cv.generated_content if isinstance(latest_cv.generated_content, dict) else {}
        cv_title = _extract_title_from_cv_markdown(content.get("markdown") or content.get("content") or "")
        if cv_title and not _is_generic_title(cv_title):
            if _extract_role_title(cv_title):
                return _normalize_title_text(f"CV {_title_case_role(cv_title)}", 72)
            return _normalize_title_text(cv_title, 72)

    return _build_chat_title_from_messages(chat_history)

async def _get_latest_cv_map(
    session: AsyncSession,
    user_id: UUID,
    conversation_ids: list[UUID],
) -> dict[UUID, GeneratedCVModel]:
    if not conversation_ids:
        return {}

    latest_versions = (
        select(
            GeneratedCVModel.conversation_id.label("conversation_id"),
            func.max(GeneratedCVModel.version).label("max_version"),
        )
        .where(
            GeneratedCVModel.user_id == user_id,
            GeneratedCVModel.conversation_id.in_(conversation_ids),
            GeneratedCVModel.deleted_at.is_(None),
        )
        .group_by(GeneratedCVModel.conversation_id)
        .subquery()
    )
    result = await session.execute(
        select(GeneratedCVModel)
        .join(
            latest_versions,
            and_(
                GeneratedCVModel.conversation_id == latest_versions.c.conversation_id,
                GeneratedCVModel.version == latest_versions.c.max_version,
            ),
        )
        .where(
            GeneratedCVModel.user_id == user_id,
            GeneratedCVModel.deleted_at.is_(None),
        )
    )
    return {item.conversation_id: item for item in result.scalars().all()}

async def _get_latest_cv_for_conversation(
    session: AsyncSession,
    user_id: UUID,
    conversation_id: UUID,
) -> GeneratedCVModel | None:
    latest_map = await _get_latest_cv_map(session, user_id, [conversation_id])
    return latest_map.get(conversation_id)

def _to_chat_session_response(
    session_model: ChatSessionModel,
    latest_cv: GeneratedCVModel | None = None,
) -> ChatSessionResponse:
    messages = session_model.messages if isinstance(session_model.messages, list) else []
    return ChatSessionResponse(
        id=session_model.id,
        conversation_id=session_model.conversation_id,
        messages=messages,
        chat_title=_build_chat_title(messages, latest_cv),
        created_at=session_model.created_at,
        updated_at=session_model.updated_at,
    )


async def _parse_uploaded_cv(file: UploadFile) -> dict[str, str]:
    settings = get_settings()
    file_bytes, upload_meta = await read_and_validate_upload(
        file,
        allowed_types={"pdf", "docx"},
        max_size_mb=min(settings.MAX_FILE_SIZE_MB, MAX_IMPORTED_CV_SIZE_MB),
        detail="Chỉ hỗ trợ file CV định dạng PDF hoặc DOCX",
    )
    normalized_name = upload_meta.filename.lower()
    ext = upload_meta.extension
    tmp_path = None
    converted_docx_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        if normalized_name.endswith(".pdf"):
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as converted:
                converted_docx_path = converted.name
            convert_pdf_to_docx(tmp_path, converted_docx_path)
            preview_payload = build_import_preview_payload(converted_docx_path)
        elif normalized_name.endswith(".docx"):
            preview_payload = build_import_preview_payload(tmp_path)
        else:
            parser = get_parser(upload_meta.filename)
            parsed_text = await parser.parse(tmp_path)
            preview_payload = {
                "markdown": parsed_text.strip(),
                "html": "",
            }
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("Failed to import uploaded CV %s: %s\n%s", upload_meta.filename, str(exc), traceback.format_exc())
        raise HTTPException(status_code=400, detail=f"Không đọc được file CV: {exc}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)
        if converted_docx_path and os.path.exists(converted_docx_path):
            os.remove(converted_docx_path)

    if not preview_payload.get("markdown", "").strip():
        raise HTTPException(status_code=400, detail="Không trích xuất được nội dung từ file CV")

    return preview_payload


def _to_generated_cv_response(cv_entity) -> GeneratedCVResponse:
    return GeneratedCVResponse(
        id=cv_entity.id,
        conversation_id=cv_entity.conversation_id,
        version=cv_entity.version,
        parent_version_id=cv_entity.parent_version_id,
        status=cv_entity.status,
        target_jd_text=cv_entity.target_jd_text,
        base_profile_data=cv_entity.base_profile_data,
        generated_content=cv_entity.generated_content,
        created_at=cv_entity.created_at,
    )


def _to_generated_cv_version_response(cv_entity) -> GeneratedCVVersionResponse:
    return GeneratedCVVersionResponse(
        id=cv_entity.id,
        conversation_id=cv_entity.conversation_id,
        version=cv_entity.version,
        parent_version_id=cv_entity.parent_version_id,
        status=cv_entity.status,
        created_at=cv_entity.created_at,
    )


def _resolve_source_info(base_profile_data: dict | None) -> tuple[str, str | None]:
    profile_data = base_profile_data if isinstance(base_profile_data, dict) else {}
    source_type = profile_data.get("source_type")
    source_filename = profile_data.get("source_filename")
    level = str(profile_data.get("level") or "").strip().lower()

    if source_type in {"uploaded_cv", "system_generated"}:
        resolved_source_type = source_type
    elif source_filename or level == "imported":
        resolved_source_type = "uploaded_cv"
    else:
        resolved_source_type = "system_generated"

    normalized_filename = str(source_filename).strip() if isinstance(source_filename, str) else None
    return resolved_source_type, normalized_filename or None


@router.post("/chat", response_model=ChatContextResponse)
async def chat_cv_generation(
    req: ChatContextRequest,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Interact with CV AI chatbot."""
    cv_repo = GeneratedCVRepository(session)
    ai_service = ai_service_factory()

    try:
        messages = [{"role": msg.role, "content": msg.content} for msg in req.messages]
        current_cv = None
        if req.current_cv_id:
            current_cv = await cv_repo.get_by_id(req.current_cv_id)
            if not current_cv or current_cv.user_id != user_id:
                raise HTTPException(status_code=404, detail="Không tìm thấy phiên bản CV hiện tại")

        if current_cv:
            use_case = EditGeneratedCVUseCase(cv_repo, ai_service)
            reply, new_cv, _ = await use_case.execute(
                user_id=user_id,
                current_cv=current_cv,
                messages=messages,
                output_format=req.output_format,
            )
            cv_id = new_cv.id if new_cv else None
        else:
            use_case = ChatCVUseCase(cv_repo, ai_service)
            reply, cv_id, conversation_id = await use_case.execute(
                user_id=user_id,
                messages=messages,
                output_format=req.output_format,
                template_id=req.template_id,
                conversation_id=req.conversation_id,
            )
            await session.commit()
            return ChatContextResponse(
                reply=reply,
                generated_cv_id=cv_id,
                conversation_id=conversation_id,
            )

        await session.commit()
            
        return ChatContextResponse(
            reply=reply,
            generated_cv_id=cv_id,
            conversation_id=current_cv.conversation_id if current_cv else req.conversation_id,
        )
    except HTTPException:
        await session.rollback()
        raise
    except Exception as e:
        await session.rollback()
        logger.error("Failed in chat interaction: %s", str(e), exc_info=True)
        raise HTTPException(status_code=500, detail="Lỗi khi giao tiếp với AI")

@router.post("/chat/stream")
async def chat_cv_generation_stream(
    req: ChatContextRequest,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Interact with CV AI chatbot via Server-Sent Events (SSE)."""
    cv_repo = GeneratedCVRepository(session)
    ai_service = ai_service_factory()

    try:
        messages = [{"role": msg.role, "content": msg.content} for msg in req.messages]
        if req.current_cv_id:
            current_cv = await cv_repo.get_by_id(req.current_cv_id)
            if not current_cv or current_cv.user_id != user_id:
                raise HTTPException(status_code=404, detail="Không tìm thấy phiên bản CV hiện tại")

            use_case = EditGeneratedCVUseCase(cv_repo, ai_service)

            async def _edit_stream():
                import json

                try:
                    yield f"event: conversation_id\ndata: {json.dumps(str(current_cv.conversation_id))}\n\n"
                    yield f"event: status\ndata: {json.dumps({'state': 'reasoning', 'label': 'AI đang phân tích yêu cầu chỉnh sửa...'})}\n\n"
                    reply, new_cv, next_content = await use_case.execute(
                        user_id=user_id,
                        current_cv=current_cv,
                        messages=messages,
                        output_format=req.output_format,
                    )
                    if reply:
                        yield f"event: chat_chunk\ndata: {json.dumps(reply)}\n\n"

                    if new_cv:
                        yield f"event: status\ndata: {json.dumps({'state': 'applying_edits', 'label': 'Đang áp thay đổi vào CV hiện tại...'})}\n\n"
                        await session.commit()
                        yield f"event: cv_chunk\ndata: {json.dumps(next_content)}\n\n"
                        yield f"event: status\ndata: {json.dumps({'state': 'saving_version', 'label': 'Đã lưu thành phiên bản CV mới.'})}\n\n"
                        yield f"event: cv_id\ndata: {json.dumps(str(new_cv.id))}\n\n"
                        yield f"event: status\ndata: {json.dumps({'state': 'done', 'label': 'Hoàn tất cập nhật CV.'})}\n\n"
                    else:
                        await session.commit()
                        yield f"event: status\ndata: {json.dumps({'state': 'waiting_input', 'label': 'Mình cần thêm thông tin trước khi sửa CV.'})}\n\n"
                except Exception as exc:
                    await session.rollback()
                    logger.error("Failed to edit CV via chat stream: %s", str(exc), exc_info=True)
                    yield f"event: error\ndata: {json.dumps(str(exc))}\n\n"

            return StreamingResponse(
                _edit_stream(),
                media_type="text/event-stream",
            )

        use_case = ChatCVUseCase(cv_repo, ai_service)
        stream_generator = use_case.execute_stream(
            user_id=user_id,
            messages=messages,
            output_format=req.output_format,
            template_id=req.template_id,
            conversation_id=req.conversation_id,
        )

        async def _stream_with_commit():
            try:
                async for chunk in stream_generator:
                    yield chunk
                await session.commit()
            except Exception:
                await session.rollback()
                raise

        return StreamingResponse(
            _stream_with_commit(),
            media_type="text/event-stream"
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Failed to start chat streaming: %s", str(e), exc_info=True)
        raise HTTPException(status_code=500, detail="Lỗi khi bắt đầu stream AI")


@router.post("/chat-sessions", response_model=ChatSessionResponse, status_code=201)
async def create_chat_session(
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Create an empty workspace chat session before a CV exists."""
    session_model = ChatSessionModel(
        user_id=user_id,
        conversation_id=uuid4(),
        messages=[],
    )
    session.add(session_model)
    await session.commit()
    await session.refresh(session_model)
    return _to_chat_session_response(session_model)


@router.get("/chat-sessions", response_model=List[ChatSessionResponse])
async def list_chat_sessions(
    limit: int = 50,
    offset: int = 0,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """List all workspace chat sessions for the current user."""
    result = await session.execute(
        select(ChatSessionModel)
        .where(
            ChatSessionModel.user_id == user_id,
        )
        .order_by(ChatSessionModel.updated_at.desc(), ChatSessionModel.created_at.desc())
        .limit(limit)
        .offset(offset)
    )
    session_models = result.scalars().all()
    latest_cv_map = await _get_latest_cv_map(
        session,
        user_id,
        [item.conversation_id for item in session_models],
    )
    return [
        _to_chat_session_response(item, latest_cv_map.get(item.conversation_id))
        for item in session_models
    ]


@router.get("/chat-sessions/{conversation_id}", response_model=ChatSessionResponse)
async def get_chat_session(
    conversation_id: UUID,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    result = await session.execute(
        select(ChatSessionModel).where(
            ChatSessionModel.conversation_id == conversation_id,
            ChatSessionModel.user_id == user_id,
        )
    )
    session_model = result.scalar_one_or_none()
    if not session_model:
        raise HTTPException(status_code=404, detail="Không tìm thấy phiên chat")
    latest_cv = await _get_latest_cv_for_conversation(session, user_id, conversation_id)
    return _to_chat_session_response(session_model, latest_cv)


@router.get("/chat-sessions/{conversation_id}/latest-cv", response_model=GeneratedCVResponse)
async def get_latest_generated_cv_by_conversation(
    conversation_id: UUID,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Get the latest active generated CV version for a conversation."""
    result = await session.execute(
        select(GeneratedCVModel)
        .where(
            GeneratedCVModel.user_id == user_id,
            GeneratedCVModel.conversation_id == conversation_id,
            GeneratedCVModel.deleted_at.is_(None),
        )
        .order_by(GeneratedCVModel.version.desc(), GeneratedCVModel.created_at.desc())
        .limit(1)
    )
    latest_cv = result.scalar_one_or_none()
    if not latest_cv:
        raise HTTPException(status_code=404, detail="Không tìm thấy CV trong cuộc trò chuyện này")

    cv_repo = GeneratedCVRepository(session)
    latest_entity = await cv_repo.get_by_id(latest_cv.id)
    return _to_generated_cv_response(latest_entity)


@router.put("/chat-sessions/{conversation_id}/messages", response_model=ChatSessionResponse)
async def update_chat_session_messages(
    conversation_id: UUID,
    req: ChatSessionMessagesRequest,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    cv_repo = GeneratedCVRepository(session)
    messages = [{"role": msg.role, "content": msg.content} for msg in req.messages]
    await cv_repo.save_chat_messages(conversation_id, user_id, messages)
    await session.commit()

    result = await session.execute(
        select(ChatSessionModel).where(
            ChatSessionModel.conversation_id == conversation_id,
            ChatSessionModel.user_id == user_id,
        )
    )
    session_model = result.scalar_one()
    latest_cv = await _get_latest_cv_for_conversation(session, user_id, conversation_id)
    return _to_chat_session_response(session_model, latest_cv)


@router.delete("/chat-sessions/{conversation_id}", status_code=204)
async def delete_chat_session(
    conversation_id: UUID,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Delete chat session by conversation id."""
    try:
        result = await session.execute(
            select(ChatSessionModel).where(
                ChatSessionModel.conversation_id == conversation_id,
                ChatSessionModel.user_id == user_id,
            )
        )
        session_model = result.scalar_one_or_none()
        if not session_model:
            raise HTTPException(status_code=404, detail="Không tìm thấy phiên chat")

        deleted_at = datetime.utcnow()
        await session.execute(
            update(GeneratedCVModel)
            .where(
                GeneratedCVModel.user_id == user_id,
                GeneratedCVModel.conversation_id == conversation_id,
                GeneratedCVModel.deleted_at.is_(None),
            )
            .values(deleted_at=deleted_at)
        )
        await session.delete(session_model)
        await session.commit()
        return None
    except HTTPException:
        await session.rollback()
        raise
    except Exception as exc:
        await session.rollback()
        logger.error(
            "Delete chat session failed: conversation_id=%s, user_id=%s, error=%s",
            conversation_id,
            user_id,
            exc,
            exc_info=True,
        )
        raise HTTPException(status_code=500, detail="Không thể xóa đoạn chat. Vui lòng thử lại.")


@router.post("/", response_model=GeneratedCVResponse)
async def generate_cv(
    req: GenerateCVRequest,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Generate a template CV via AI and save it."""
    
    cv_repo = GeneratedCVRepository(session)
    user_repo = UserRepository(session)
    ai_service = ai_service_factory()

    user = await user_repo.get_by_id(user_id)
    user_profile = None
    if user is not None:
        user_profile = {
            "full_name": user.full_name or "",
            "email": user.email or "",
            "phone_number": user.phone_number or "",
        }

    use_case = GenerateCVUseCase(cv_repo, ai_service)

    try:
        cv_entity = await use_case.execute(
            user_id=user_id,
            job_title=req.job_title,
            jd_text=req.jd_text,
            level=req.level,
            output_format=req.output_format,
            user_profile=user_profile,
        )
        await session.commit()
        return _to_generated_cv_response(cv_entity)

    except HTTPException:
        await session.rollback()
        raise
    except Exception as e:
        await session.rollback()
        logger.error("Failed to generate CV: %s", str(e), exc_info=True)
        raise HTTPException(status_code=500, detail="Lỗi khi AI tạo CV mẫu")


@router.post("/import", response_model=GeneratedCVResponse, status_code=201)
async def import_generated_cv(
    cv_file: UploadFile = File(...),
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Import an existing CV file and open it in the editable workspace."""
    cv_repo = GeneratedCVRepository(session)
    use_case = ImportGeneratedCVUseCase(cv_repo)

    logger.info(
        "Import generated CV request: user_id=%s, filename=%s, content_type=%s",
        user_id,
        cv_file.filename,
        cv_file.content_type,
    )

    try:
        preview_payload = await _parse_uploaded_cv(cv_file)
        cv_entity = await use_case.execute(
            user_id=user_id,
            filename=cv_file.filename or "uploaded_cv",
            parsed_content=preview_payload["markdown"],
            preview_html=preview_payload.get("html", ""),
        )
        await session.commit()
        logger.info("Imported CV saved successfully: cv_id=%s, filename=%s", cv_entity.id, cv_file.filename)
        return _to_generated_cv_response(cv_entity)
    except HTTPException:
        await session.rollback()
        raise
    except ValueError as exc:
        await session.rollback()
        raise HTTPException(status_code=400, detail=str(exc))
    except Exception as exc:
        await session.rollback()
        logger.error("Failed to import CV into workspace: %s", str(exc), exc_info=True)
        raise HTTPException(status_code=500, detail="Không thể import CV vào workspace")


@router.post("/{cv_id}/import-version", response_model=GeneratedCVResponse, status_code=201)
async def import_generated_cv_version(
    cv_id: UUID,
    cv_file: UploadFile = File(...),
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Import an uploaded CV as a new version of the current workspace conversation."""
    cv_repo = GeneratedCVRepository(session)
    cv_entity = await cv_repo.get_by_id(cv_id)

    if not cv_entity or cv_entity.user_id != user_id:
        raise HTTPException(status_code=404, detail="Không tìm thấy CV hiện tại")

    try:
        preview_payload = await _parse_uploaded_cv(cv_file)
        existing_payload = cv_entity.generated_content if isinstance(cv_entity.generated_content, dict) else {}
        existing_history = existing_payload.get("chat_history") if isinstance(existing_payload.get("chat_history"), list) else []
        filename = cv_file.filename or "uploaded_cv"
        imported_message = (
            f"Đã nhận diện tài liệu đính kèm là CV và thay vào workspace hiện tại từ file `{filename}`. "
            "Bạn có thể tiếp tục chỉnh sửa trên bản này."
        )
        content = preview_payload["markdown"].strip()
        preview_html = preview_payload.get("html", "").strip()

        new_entity = await cv_repo.create_versioned(
            user_id=user_id,
            conversation_id=cv_entity.conversation_id,
            parent_version_id=cv_entity.id,
            target_jd_text=cv_entity.target_jd_text,
            base_profile_data={
                **(cv_entity.base_profile_data or {}),
                "source_filename": filename,
                "source_type": "uploaded_cv",
            },
            generated_content={
                "format": "markdown",
                "content": content,
                "markdown": content,
                "html": preview_html,
                "import_preview_format": "html" if preview_html else "markdown",
                "source_filename": filename,
                "chat_history": [
                    *existing_history,
                    {"role": "assistant", "content": imported_message},
                ],
            },
            status=cv_entity.status,
        )
        await session.commit()
        logger.info("Imported CV as workspace version: old_cv_id=%s, new_cv_id=%s", cv_id, new_entity.id)
        return _to_generated_cv_response(new_entity)
    except HTTPException:
        await session.rollback()
        raise
    except Exception as exc:
        await session.rollback()
        logger.error("Failed to import CV as workspace version: %s", str(exc), exc_info=True)
        raise HTTPException(status_code=500, detail="Không thể thay CV đính kèm vào workspace hiện tại")


@router.get("/", response_model=List[GeneratedCVListResponse])
async def list_generated_cvs(
    limit: int = 20,
    offset: int = 0,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """List user's generated CVs."""
    cv_repo = GeneratedCVRepository(session)
    cvs = await cv_repo.list_by_user_id(user_id, limit, offset)

    items = []
    for c in cvs:
        cv_with_chat = await cv_repo.get_by_id(c.id)
        content_data = cv_with_chat.generated_content if cv_with_chat else c.generated_content
        chat_title = _build_chat_title((content_data or {}).get("chat_history"), c)
        source_type, source_filename = _resolve_source_info(c.base_profile_data)
        preview_markdown = ""
        preview_html = ""
        preview_format = None
        if source_type == "uploaded_cv" and isinstance(content_data, dict):
            preview_markdown = str(content_data.get("markdown") or content_data.get("content") or "").strip()
            preview_html = str(content_data.get("html") or "").strip()
            preview_format = str(content_data.get("import_preview_format") or ("html" if preview_html else "markdown")).strip()

        items.append(GeneratedCVListResponse(
            id=c.id,
            conversation_id=c.conversation_id,
            version=c.version,
            status=c.status,
            chat_title=chat_title,
            target_jd_text=c.target_jd_text,
            job_title=c.base_profile_data.get("job_title") if c.base_profile_data else None,
            level=c.base_profile_data.get("level") if c.base_profile_data else None,
            source_type=source_type,
            source_filename=source_filename,
            is_editable=True,
            preview_markdown=preview_markdown[:6000] if preview_markdown else None,
            preview_html=preview_html[:6000] if preview_html else None,
            preview_format=preview_format or None,
            created_at=c.created_at,
        ))
    return items


@router.get("/{cv_id}", response_model=GeneratedCVResponse)
async def get_generated_cv(
    cv_id: UUID,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Get full generated CV."""
    cv_repo = GeneratedCVRepository(session)
    cv_entity = await cv_repo.get_by_id(cv_id)
    
    if not cv_entity or cv_entity.user_id != user_id:
        raise HTTPException(status_code=404, detail="Không tìm thấy CV mẫu này")

    return _to_generated_cv_response(cv_entity)


@router.get("/{cv_id}/versions", response_model=List[GeneratedCVVersionResponse])
async def list_generated_cv_versions(
    cv_id: UUID,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    cv_repo = GeneratedCVRepository(session)
    cv_entity = await cv_repo.get_by_id(cv_id)
    if not cv_entity or cv_entity.user_id != user_id:
        raise HTTPException(status_code=404, detail="Không tìm thấy CV mẫu này")

    versions = await cv_repo.list_versions(user_id, cv_entity.conversation_id)
    return [_to_generated_cv_version_response(item) for item in versions]


async def _download_generated_cv(
    cv_id: UUID,
    format: Literal["markdown", "docx"] | None = Query(default=None),
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Download generated CV as markdown or docx."""
    cv_repo = GeneratedCVRepository(session)
    cv_entity = await cv_repo.get_by_id(cv_id)

    if not cv_entity or cv_entity.user_id != user_id:
        raise HTTPException(status_code=404, detail="Không tìm thấy CV mẫu này")

    stored_format, stored_content = _get_generated_content_payload(cv_entity)
    content_data = cv_entity.generated_content if isinstance(cv_entity.generated_content, dict) else {}
    export_format = format or stored_format

    if export_format not in {"markdown", "docx"}:
        raise HTTPException(status_code=400, detail="Định dạng export không hợp lệ")

    export_content = content_data.get("markdown") or stored_content

    if not str(export_content).strip():
        raise HTTPException(status_code=400, detail="CV không có nội dung để export")

    if export_format == "docx":
        docx_content = _markdown_to_docx_bytes(str(export_content))
        filename = _build_export_filename(cv_entity, "docx")
        return StreamingResponse(
            io.BytesIO(docx_content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    if export_format == "markdown":
        filename = _build_export_filename(cv_entity, "md")
        return Response(
            content=str(export_content),
            media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )


@router.get("/{cv_id}/download")
async def download_generated_cv(
    cv_id: UUID,
    format: Literal["markdown", "docx"] | None = Query(default=None),
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    return await _download_generated_cv(cv_id, format, user_id, session)


@router.get("/{cv_id}/export")
async def export_generated_cv(
    cv_id: UUID,
    format: Literal["markdown", "docx"] | None = Query(default=None),
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    return await _download_generated_cv(cv_id, format, user_id, session)


@router.post("/export-preview-docx")
async def export_preview_docx(
    req: GeneratedCVExportPreviewRequest,
    user_id: UUID = Depends(get_current_user_id),
):
    """Export current editor markdown content to DOCX without requiring a saved CV."""
    markdown_content = str(req.content or "").strip()
    if not markdown_content:
        raise HTTPException(status_code=400, detail="CV không có nội dung để export")

    docx_content = _markdown_to_docx_bytes(markdown_content)
    fallback_title = str(req.title or "").strip() or "generated_cv"
    normalized = re.sub(r"[^a-zA-Z0-9_-]+", "_", fallback_title.lower()).strip("_")
    filename = f"{(normalized or 'generated_cv')[:60]}.docx"
    return StreamingResponse(
        io.BytesIO(docx_content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/{cv_id}/versions", response_model=GeneratedCVResponse)
async def create_generated_cv_version(
    cv_id: UUID,
    req: GeneratedCVUpdateRequest,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Create a new immutable version after user edits in preview."""
    cv_repo = GeneratedCVRepository(session)
    cv_entity = await cv_repo.get_by_id(cv_id)

    if not cv_entity or cv_entity.user_id != user_id:
        raise HTTPException(status_code=404, detail="Không tìm thấy CV mẫu này")

    existing_payload = cv_entity.generated_content if isinstance(cv_entity.generated_content, dict) else {}
    next_payload = {
        key: value
        for key, value in existing_payload.items()
        if key not in {"html", "import_preview_format"}
    }
    new_entity = await cv_repo.create_versioned(
        user_id=user_id,
        conversation_id=cv_entity.conversation_id,
        parent_version_id=cv_entity.id,
        target_jd_text=cv_entity.target_jd_text,
        base_profile_data=cv_entity.base_profile_data,
        generated_content={
            **next_payload,
            "format": req.output_format,
            "content": req.content,
            "markdown": req.content,
        },
        status=cv_entity.status,
    )

    await session.commit()
    return _to_generated_cv_response(new_entity)


@router.patch("/{cv_id}", response_model=GeneratedCVResponse)
async def update_generated_cv(
    cv_id: UUID,
    req: GeneratedCVUpdateRequest,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    return await create_generated_cv_version(cv_id, req, user_id, session)

@router.delete("/{cv_id}", status_code=204)
async def delete_generated_cv(
    cv_id: UUID,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Soft delete generated CV."""
    cv_repo = GeneratedCVRepository(session)
    success = await cv_repo.soft_delete(cv_id, user_id)
    
    if not success:
        raise HTTPException(status_code=404, detail="Không tìm thấy CV mẫu này")

    await session.commit()
