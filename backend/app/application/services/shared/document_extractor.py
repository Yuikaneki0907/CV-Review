"""Single document → plain-text extractor.

Replaces the four duplicated tempfile + parser dances scattered across
``presentation/analysis_routes.py`` and ``presentation/generated_cv_routes.py``.
Operates on in-memory bytes — no caller needs to write a temp file
unless an external library demands a path, and even then this helper
hides it.

Phase 1 will switch route handlers to call this. The legacy
``infrastructure/file_parsers/parsers.py`` stays around because
``import_pipeline`` (PDF → DOCX → HTML) is a separate concern.
"""
from __future__ import annotations

import io
import os
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

from app.application.exceptions import DocumentExtractionError
from app.domain.schemas import ExtractedDocument
from app.domain.schemas.document_schema import DocumentExtension
from app.logger import get_logger

logger = get_logger("app.application.services.shared.document_extractor")

# Below this length we flag the extraction as low-quality. Most scanned
# PDFs and corrupt uploads produce <200 characters of usable text.
_LOW_QUALITY_CHAR_THRESHOLD = 200


def _detect_extension(filename: str) -> DocumentExtension:
    """Map a filename to a normalised extension."""
    lowered = (filename or "").lower()
    suffix = Path(lowered).suffix.lstrip(".")
    if suffix == "pdf":
        return "pdf"
    if suffix == "docx":
        return "docx"
    if suffix == "txt":
        return "txt"
    if suffix == "md":
        return "md"
    raise DocumentExtractionError(f"Unsupported file extension: {filename!r}")


def _extract_pdf(file_bytes: bytes) -> tuple[str, int]:
    """Return (text, page_count) for a PDF byte stream."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:  # pymupdf raises bare Exception subclasses
        raise DocumentExtractionError(f"Failed to open PDF: {exc}") from exc
    try:
        pages = [page.get_text() for page in doc]
        page_count = doc.page_count
    finally:
        doc.close()
    return "\n".join(pages), page_count


def _extract_docx(file_bytes: bytes) -> str:
    """Return text for a DOCX byte stream.

    python-docx demands a file path or a file-like object — we use
    ``BytesIO`` to avoid touching disk.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise DocumentExtractionError(f"Failed to open DOCX: {exc}") from exc
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_plaintext(file_bytes: bytes) -> str:
    """Decode a UTF-8 text/markdown byte stream."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        # Fall back to latin-1 — every byte sequence decodes, so we get
        # *something* and flag it as low-quality downstream.
        return file_bytes.decode("latin-1", errors="replace")


async def extract_document_text(
    file_bytes: bytes,
    filename: str,
) -> ExtractedDocument:
    """Parse an uploaded document into plain text.

    Args:
        file_bytes: Raw file contents.
        filename: Original filename — used only to detect extension.

    Returns:
        :class:`ExtractedDocument` with the parsed text, extension,
        page count (PDF only), an ``extraction_quality`` flag, and
        any warnings the caller may want to surface.

    Raises:
        DocumentExtractionError: When the format is unsupported or the
            file cannot be parsed at all. Callers should map this to a
            400-level HTTP response.
    """
    extension = _detect_extension(filename)
    page_count: int | None = None

    if extension == "pdf":
        text, page_count = _extract_pdf(file_bytes)
    elif extension == "docx":
        text = _extract_docx(file_bytes)
    else:
        text = _extract_plaintext(file_bytes)

    warnings: list[str] = []
    quality = "high"
    stripped_len = len(text.strip())
    if stripped_len < _LOW_QUALITY_CHAR_THRESHOLD:
        quality = "low"
        warnings.append(
            f"{extension} yielded {stripped_len} chars (<{_LOW_QUALITY_CHAR_THRESHOLD})"
        )

    logger.info(
        "extract_document_text: filename=%s ext=%s chars=%d pages=%s quality=%s",
        filename,
        extension,
        stripped_len,
        page_count,
        quality,
    )

    return ExtractedDocument(
        text=text,
        extension=extension,
        page_count=page_count,
        extraction_quality=quality,
        warnings=warnings,
    )


def write_to_tempfile(file_bytes: bytes, extension: str) -> str:
    """Escape hatch for libraries that demand a path.

    Returns the absolute path of a new tempfile that the caller MUST
    delete. Prefer :func:`extract_document_text` whenever possible.
    """
    suffix = extension if extension.startswith(".") else f".{extension}"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(file_bytes)
        return tmp.name


def remove_tempfile(path: str | None) -> None:
    """Best-effort tempfile cleanup."""
    if path and os.path.exists(path):
        try:
            os.remove(path)
        except OSError:
            logger.warning("Failed to remove tempfile: %s", path, exc_info=True)
